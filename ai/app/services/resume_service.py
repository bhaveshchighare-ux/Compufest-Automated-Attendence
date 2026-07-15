import uuid
import os
import tempfile
from datetime import datetime
from sqlalchemy import func

from app.persistence.database.session import get_db_session
from app.persistence.database.models.user_model import UserModel
from app.persistence.database.models.resume_model import Resume
from app.persistence.redis.resume_store import ResumeDataStore
from app.core.exceptions import ValidationError, NotFoundError
from app.core.logger import get_logger

import pdfplumber
from docx import Document
import fitz  # PyMuPDF for PDF image extraction
import os
from pathlib import Path

logger = get_logger(__name__)



class ResumeService:

    def create_resume_from_file(self, file, name: str, email: str, candidate_image_file=None, candidate_image_vector: str = None) -> dict:
        # Extract text BEFORE opening database session to prevent long-held connections
        extracted_text = self._extract_text(file)

        logger.info(f"Extracted text length: {len(extracted_text)} characters")
        logger.debug(f"Extracted text preview: {extracted_text[:200] if extracted_text else 'NO TEXT'}")

        # Allow empty text (e.g., for scanned PDFs) - continue with filename as fallback
        if not extracted_text.strip():
            logger.warning(f"No text extracted from {file.filename}, using filename as description")
            extracted_text = f"Resume: {file.filename}"

        # Under pure vector architecture, we bypass all server-side extraction and file writing.
        # Only the 128-d reference vector string from the client is saved.
        face_image_path = candidate_image_vector

        try:
            db = get_db_session()
            
            user_exist = db.query(UserModel).filter(func.lower(UserModel.email) == email.lower()).first()

            if user_exist:
                candidate_id = user_exist.id
                # Keep user profile aligned with latest submitted name.
                if name and user_exist.name != name:
                    user_exist.name = name
            else:
                candidate_id = str(uuid.uuid4())
                user = UserModel(
                    id=candidate_id,
                    name=name,
                    email=email.lower(),
                )
                db.add(user)
                db.flush()

            # UPSERT: Update existing resume if one exists, otherwise create new
            existing_resume = db.query(Resume).filter(Resume.candidate_id == candidate_id).first()

            if existing_resume:
                # Update existing resume in-place
                existing_resume.filename = file.filename
                existing_resume.extracted_text = extracted_text
                existing_resume.face_image_path = face_image_path
                existing_resume.updated_at = datetime.utcnow()
                resume_id = str(existing_resume.id)
                logger.info(f"Updated existing resume {resume_id} for candidate {candidate_id}")
            else:
                # First-time upload: create new resume row
                resume_id = str(uuid.uuid4())
                resume = Resume(
                    id=resume_id,
                    candidate_id=candidate_id,
                    filename=file.filename,
                    extracted_text=extracted_text,
                    face_image_path=face_image_path,
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow(),
                )
                db.add(resume)
                logger.info(f"Created new resume {resume_id} for candidate {candidate_id}")

            db.commit()
            
            # Cache the newly created/updated resume data in Redis
            ResumeDataStore().set_resume_data(
                resume_id=resume_id,
                text=extracted_text,
                face_vector=face_image_path
            )
        except Exception as e:
            #Ensure session is closed even on error
            if 'db' in locals():
                db.close()
            raise
        finally:
            if 'db' in locals():
                db.close()

        return {
            "candidate_id": candidate_id,
            "resume_id": resume_id,
            "has_face_image": bool(face_image_path),
        }

    def scan_resume_for_face(self, file) -> dict:
        # Enforce direct client-side photo/vector capture for highly accurate and secure BIPA/GDPR compliant proctoring.
        return {"has_face_image": False}

    def attach_reference_image(self, resume_id: str, image_file) -> dict:
        if image_file is None or not getattr(image_file, "filename", ""):
            raise ValidationError("Candidate image is required")

        from werkzeug.utils import secure_filename
        allowed_extensions = {"jpg", "jpeg", "png", "webp"}
        filename = secure_filename(image_file.filename)
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in allowed_extensions:
            raise ValidationError("Unsupported image format. Allowed: jpg, jpeg, png, webp")

        db = get_db_session()
        try:
            resume = db.query(Resume).filter(Resume.id == resume_id).first()
            if not resume:
                raise NotFoundError("Resume not found")

            image_path = self._store_candidate_image(resume_id, image_file)
            resume.face_image_path = str(image_path)
            resume.updated_at = datetime.utcnow()
            db.commit()
            
            # Update Redis cache
            ResumeDataStore().set_resume_data(
                resume_id=resume_id,
                text=resume.extracted_text,
                face_vector=str(image_path)
            )

            return {
                "resume_id": resume_id,
                "face_image_path": str(image_path),
                "has_face_image": True,
            }
        finally:
            db.close()

    def _extract_text(self, file) -> str:
        filename = file.filename.lower()

        if filename.endswith(".pdf"):
            return self._extract_pdf(file)

        if filename.endswith(".docx"):
            return self._extract_docx(file)

        raise ValidationError("Unsupported resume format")

    def _extract_pdf(self, file) -> str:
        logger.info(f"Extracting PDF: {file.filename}")
        text = []
        try:
            with pdfplumber.open(file) as pdf:
                logger.info(f"PDF has {len(pdf.pages)} pages")
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                    logger.debug(f"Extracted page {i+1}/{len(pdf.pages)}")
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValidationError(f"Failed to extract PDF: {str(e)}")
        logger.info("PDF extraction completed")
        return "\n".join(text)

    def _extract_docx(self, file) -> str:
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)

    def _extract_image_from_pdf(self, file_path: str, resume_id: str) -> str:
        """Extract face image from PDF resume"""
        try:
            doc = fitz.open(file_path)
            
            # Strategy 1: Try to extract embedded image objects
            for page_index in range(len(doc)):
                page = doc[page_index]
                images = page.get_images(full=True)
                
                for img in images:
                    xref = img[0]
                    base = doc.extract_image(xref)
                    image_bytes = base["image"]
                    
                    # Try to detect and crop a face from this specific image
                    face_data = self._detect_and_crop_face(image_bytes)
                    if face_data:
                        uploads_dir = Path("uploads")
                        uploads_dir.mkdir(exist_ok=True)
                        img_path = uploads_dir / f"{resume_id}_face.jpg"
                        with open(img_path, "wb") as f:
                            f.write(face_data)
                        return str(img_path)
            
            # Strategy 2: Fallback - Render pages and scan for a face
            # Scan first 2 pages (usually enough for photos)
            for page_index in range(min(2, len(doc))):
                page = doc[page_index]
                # Render page at high resolution
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                page_img_bytes = pix.tobytes("jpg")
                
                face_data = self._detect_and_crop_face(page_img_bytes)
                if face_data:
                    uploads_dir = Path("uploads")
                    uploads_dir.mkdir(exist_ok=True)
                    img_path = uploads_dir / f"{resume_id}_face.jpg"
                    with open(img_path, "wb") as f:
                        f.write(face_data)
                    return str(img_path)
            
            return None
        except Exception as e:
            logger.warning(f"Image extraction failed: {e}")
            return None

    def _pdf_contains_face_image(self, file_path: str) -> bool:
        try:
            doc = fitz.open(file_path)

            # Check embedded images
            for page_index in range(len(doc)):
                page = doc[page_index]
                images = page.get_images(full=True)

                for img in images:
                    xref = img[0]
                    base = doc.extract_image(xref)
                    if self._image_contains_face(base["image"]) or self._is_substantial_pdf_image(base):
                        return True
            
            # Check full page render as fallback (first 2 pages)
            for i in range(min(2, len(doc))):
                pix = doc[i].get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                if self._image_contains_face(pix.tobytes("jpg")):
                    return True

            return False
        except Exception as e:
            logger.warning(f"Resume face scan failed: {e}")
            return False

    def _store_candidate_image(self, resume_id: str, image_file) -> str:
        from werkzeug.utils import secure_filename
        allowed_extensions = {"jpg", "jpeg", "png", "webp"}
        filename = secure_filename(image_file.filename)
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in allowed_extensions:
            raise ValidationError("Unsupported image format. Allowed: jpg, jpeg, png, webp")

        image_file.seek(0)
        image_bytes = image_file.read()
        if not self._image_contains_face(image_bytes):
            raise ValidationError("Uploaded image must contain one clear front-facing face")

        uploads_dir = Path("uploads")
        uploads_dir.mkdir(exist_ok=True)
        image_path = uploads_dir / f"{resume_id}_face.{ext}"
        with open(image_path, "wb") as image_output:
            image_output.write(image_bytes)
        image_file.seek(0)
        return str(image_path)

    def _docx_contains_face_image(self, file) -> bool:
        try:
            doc = Document(file)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_bytes = rel.target_part.blob
                    if self._image_contains_face(image_bytes):
                        return True
            return False
        except Exception as e:
            logger.warning(f"DOCX face scan failed: {e}")
            return False

    def _extract_image_from_docx(self, file, resume_id: str) -> str:
        try:
            doc = Document(file)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_bytes = rel.target_part.blob
                    face_data = self._detect_and_crop_face(image_bytes)
                    if face_data:
                        uploads_dir = Path("uploads")
                        uploads_dir.mkdir(exist_ok=True)
                        img_path = uploads_dir / f"{resume_id}_face.jpg"
                        with open(img_path, "wb") as f:
                            f.write(face_data)
                        return str(img_path)
            return None
        except Exception as e:
            logger.warning(f"DOCX image extraction failed: {e}")
            return None

    def _image_contains_face(self, image_bytes: bytes) -> bool:
        return self._detect_and_crop_face(image_bytes) is not None

    def _detect_and_crop_face(self, image_bytes: bytes) -> bytes:
        """
        Detects a face in an image using a multi-stage pipeline:
        1. Preprocessing (CLAHE) for better contrast.
        2. Fast Haar Cascade check.
        3. High-precision DeepFace fallback.
        """
        import cv2
        import numpy as np
        from deepface import DeepFace

        nparr = np.frombuffer(image_bytes, np.uint8)
        img_cv = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if img_cv is None:
            return None

        # --- Phase 1: Preprocessing ---
        # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization) to improve detection in poor lighting
        lab = cv2.cvtColor(img_cv, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
        cl = clahe.apply(l)
        limg = cv2.merge((cl,a,b))
        img_enhanced = cv2.cvtColor(limg, cv2.COLOR_LAB2BGR)
        gray = cv2.cvtColor(img_enhanced, cv2.COLOR_BGR2GRAY)

        # --- Phase 2: Fast Haar Cascades ---
        cascades = ['haarcascade_frontalface_default.xml', 'haarcascade_frontalface_alt.xml']
        for cascade_name in cascades:
            cascade_path = cv2.data.haarcascades + cascade_name
            face_cascade = cv2.CascadeClassifier(cascade_path)
            if face_cascade.empty(): continue
            
            faces = face_cascade.detectMultiScale(gray, 1.05, 3, minSize=(30, 30))
            if len(faces) > 0:
                (x, y, w, h) = max(faces, key=lambda f: f[2] * f[3])
                # Add padding for a nice crop
                pad = int(h * 0.2)
                y1, y2 = max(0, y - pad), min(img_cv.shape[0], y + h + pad)
                x1, x2 = max(0, x - pad), min(img_cv.shape[1], x + w + pad)
                face_img = img_cv[y1:y2, x1:x2]
                _, buffer = cv2.imencode(".jpg", face_img)
                return buffer.tobytes()

        # --- Phase 3: High-Precision DeepFace Fallback ---
        try:
            # Using retinaface or mtcnn as backends (retinaface is most accurate but slower)
            # detect_backend can be 'opencv', 'retinaface', 'mtcnn', 'ssd', 'dlib', 'mediapipe'
            objs = DeepFace.extract_faces(
                img_path=img_cv, 
                detector_backend='opencv', # Try opencv backend first in DeepFace too
                enforce_detection=False,
                align=True
            )
            
            # If opencv failed, try a stronger one
            if not objs or objs[0]['confidence'] < 0.8:
                 objs = DeepFace.extract_faces(
                    img_path=img_cv, 
                    detector_backend='mtcnn', 
                    enforce_detection=False,
                    align=True
                )

            if objs and len(objs) > 0 and objs[0]['confidence'] > 0.4:
                face_obj = objs[0]
                face_img = (face_obj['face'] * 255).astype(np.uint8)
                face_img = cv2.cvtColor(face_img, cv2.COLOR_RGB2BGR)
                _, buffer = cv2.imencode(".jpg", face_img)
                return buffer.tobytes()
        except Exception as e:
            logger.debug(f"DeepFace extraction fallback failed: {e}")

        return None

    def _is_substantial_pdf_image(self, image_info: dict) -> bool:
        width = int(image_info.get("width") or 0)
        height = int(image_info.get("height") or 0)
        image_size = len(image_info.get("image") or b"")
        return width >= 60 and height >= 60 and image_size >= 2000

    def check_existing_resume(self, email: str) -> dict:
        """Check if a resume already exists for the given email."""
        db = get_db_session()
        try:
            user = db.query(UserModel).filter(func.lower(UserModel.email) == email.lower()).first()
            if not user:
                return {"has_resume": False}

            resume = db.query(Resume).filter(Resume.candidate_id == user.id).first()
            if not resume:
                return {"has_resume": False}

            return {
                "has_resume": True,
                "candidate_id": str(resume.candidate_id),
                "resume_id": str(resume.id),
                "filename": resume.filename,
                "updated_at": resume.updated_at.isoformat() if resume.updated_at else None,
            }
        finally:
            db.close()

    def clear_face_vector(self, resume_id: str) -> None:
        """Wipe face_image_path after interview ends for privacy."""
        db = get_db_session()
        try:
            resume = db.query(Resume).filter(Resume.id == resume_id).first()
            if resume and resume.face_image_path:
                resume.face_image_path = None
                resume.updated_at = datetime.utcnow()
                db.commit()
                logger.info(f"Cleared face vector for resume {resume_id}")
            
            # Clear from Redis cache as well
            ResumeDataStore().clear_face_vector(resume_id)
        except Exception as e:
            logger.warning(f"Failed to clear face vector for resume {resume_id}: {e}")
        finally:
            db.close()

    def get_resume_text(self, resume_id: str) -> str:
        # Try fetching from Redis first
        cached_data = ResumeDataStore().get_resume_data(resume_id)
        if cached_data and cached_data.get("text"):
            return cached_data["text"]

        db = get_db_session()
        try:
            resume = db.query(Resume).filter(Resume.id == resume_id).first()

            if not resume:
                raise NotFoundError("Resume not found")

            if not resume.extracted_text:
                raise ValidationError("Resume text not available")

            # Populate cache
            ResumeDataStore().set_resume_data(
                resume_id=resume_id,
                text=resume.extracted_text,
                face_vector=resume.face_image_path
            )

            return resume.extracted_text
        finally:
            db.close()

    def get_resume_face_image(self, resume_id: str) -> str:
        # Try fetching from Redis first
        cached_data = ResumeDataStore().get_resume_data(resume_id)
        if cached_data and "face_vector" in cached_data:
            return cached_data["face_vector"] or None

        db = get_db_session()
        try:
            resume = db.query(Resume).filter(Resume.id == resume_id).first()

            if not resume:
                raise NotFoundError("Resume not found")

            # Populate cache
            ResumeDataStore().set_resume_data(
                resume_id=resume_id,
                text=resume.extracted_text,
                face_vector=resume.face_image_path
            )

            return resume.face_image_path
        finally:
            db.close()

